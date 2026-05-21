"""
文件处理服务
支持多模态文件解析：图片、PPT、PDF、Excel、音频、视频
"""
import os
import io
from typing import Dict, Any, Optional
import pandas as pd
from PyPDF2 import PdfReader
from docx import Document
from PIL import Image
import moviepy.editor as mp
import speech_recognition as sr
from pydub import AudioSegment
import logging

logger = logging.getLogger(__name__)


class FileProcessor:
    """文件处理器"""
    
    @staticmethod
    async def process_file(file_path: str, file_type: str) -> Dict[str, Any]:
        """
        处理文件，根据类型调用不同的解析器
        
        Args:
            file_path: 文件路径
            file_type: 文件类型（excel/csv/pdf/word/image/audio/video）
        
        Returns:
            解析结果
        """
        try:
            if file_type in ["excel", "csv"]:
                return await FileProcessor._parse_excel_csv(file_path, file_type)
            elif file_type == "pdf":
                return await FileProcessor._parse_pdf(file_path)
            elif file_type == "word":
                return await FileProcessor._parse_word(file_path)
            elif file_type == "image":
                return await FileProcessor._parse_image(file_path)
            elif file_type == "audio":
                return await FileProcessor._parse_audio(file_path)
            elif file_type == "video":
                return await FileProcessor._parse_video(file_path)
            else:
                return {"error": f"不支持的文件类型：{file_type}"}
        except Exception as e:
            logger.error(f"文件处理失败：{str(e)}")
            return {"error": str(e)}
    
    @staticmethod
    async def _parse_excel_csv(file_path: str, file_type: str) -> Dict[str, Any]:
        """解析 Excel/CSV 文件"""
        try:
            if file_type == "excel":
                df = pd.read_excel(file_path)
            else:
                df = pd.read_csv(file_path)
            
            return {
                "type": "table",
                "rows": len(df),
                "columns": list(df.columns),
                "data": df.to_dict(orient="records"),
                "preview": df.head(10).to_dict(orient="records")
            }
        except Exception as e:
            return {"error": f"Excel/CSV 解析失败：{str(e)}"}
    
    @staticmethod
    async def _parse_pdf(file_path: str) -> Dict[str, Any]:
        """解析 PDF 文件"""
        try:
            reader = PdfReader(file_path)
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
            
            return {
                "type": "text",
                "pages": len(reader.pages),
                "text": text[:5000],  # 限制长度
                "preview": text[:1000]
            }
        except Exception as e:
            return {"error": f"PDF 解析失败：{str(e)}"}
    
    @staticmethod
    async def _parse_word(file_path: str) -> Dict[str, Any]:
        """解析 Word 文件"""
        try:
            doc = Document(file_path)
            text = "\n".join([para.text for para in doc.paragraphs])
            
            return {
                "type": "text",
                "paragraphs": len(doc.paragraphs),
                "text": text[:5000],
                "preview": text[:1000]
            }
        except Exception as e:
            return {"error": f"Word 解析失败：{str(e)}"}
    
    @staticmethod
    async def _parse_image(file_path: str) -> Dict[str, Any]:
        """解析图片文件"""
        try:
            img = Image.open(file_path)
            width, height = img.size
            
            return {
                "type": "image",
                "width": width,
                "height": height,
                "format": img.format,
                "mode": img.mode
            }
        except Exception as e:
            return {"error": f"图片解析失败：{str(e)}"}
    
    @staticmethod
    async def _parse_audio(file_path: str) -> Dict[str, Any]:
        """解析音频文件"""
        try:
            # 转换为 WAV 格式
            audio = AudioSegment.from_file(file_path)
            audio.export("temp.wav", format="wav")
            
            # 语音识别
            recognizer = sr.Recognizer()
            with sr.AudioFile("temp.wav") as source:
                audio_data = recognizer.record(source)
                text = recognizer.recognize_google(audio_data, language="zh-CN")
            
            return {
                "type": "audio",
                "duration": len(audio) / 1000.0,  # 秒
                "text": text,
                "preview": text[:500]
            }
        except Exception as e:
            return {"error": f"音频解析失败：{str(e)}"}
    
    @staticmethod
    async def _parse_video(file_path: str) -> Dict[str, Any]:
        """解析视频文件"""
        try:
            video = mp.VideoFileClip(file_path)
            duration = video.duration
            
            # 提取音频
            video.audio.write_audiofile("temp_audio.wav")
            
            # 解析音频
            audio_result = await FileProcessor._parse_audio("temp_audio.wav")
            
            return {
                "type": "video",
                "duration": duration,
                "fps": video.fps,
                "size": (video.w, video.h),
                "audio_text": audio_result.get("text", ""),
                "preview": audio_result.get("preview", "")
            }
        except Exception as e:
            return {"error": f"视频解析失败：{str(e)}"}


# 文件类型检测
def detect_file_type(file_path: str) -> str:
    """
    检测文件类型
    
    Returns:
        excel/csv/pdf/word/image/audio/video/unknown
    """
    ext = file_path.split(".")[-1].lower() if "." in file_path else ""
    
    if ext in ["xlsx", "xls"]:
        return "excel"
    elif ext == "csv":
        return "csv"
    elif ext == "pdf":
        return "pdf"
    elif ext in ["docx", "doc"]:
        return "word"
    elif ext in ["jpg", "jpeg", "png", "gif", "bmp", "webp"]:
        return "image"
    elif ext in ["mp3", "wav", "ogg", "flac", "aac"]:
        return "audio"
    elif ext in ["mp4", "avi", "mov", "mkv", "wmv"]:
        return "video"
    else:
        return "unknown"
