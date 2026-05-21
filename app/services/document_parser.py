"""
文档解析服务 - 支持 PDF/Word/Excel 真实解析
"""
import os
import io
import logging
from typing import Dict, Any, Optional
import pandas as pd
from PyPDF2 import PdfReader
from docx import Document
from openpyxl import load_workbook

logger = logging.getLogger(__name__)


class DocumentParser:
    """文档解析器"""
    
    @staticmethod
    async def parse_file(file_path: str, file_type: str) -> Dict[str, Any]:
        """
        解析文件
        
        Args:
            file_path: 文件路径
            file_type: 文件类型
        
        Returns:
            解析结果
        """
        try:
            if file_type == "pdf":
                return await DocumentParser._parse_pdf(file_path)
            elif file_type in ["doc", "docx"]:
                return await DocumentParser._parse_docx(file_path)
            elif file_type in ["xls", "xlsx", "csv"]:
                return await DocumentParser._parse_excel(file_path)
            else:
                return {"error": f"不支持的文件类型：{file_type}"}
        except Exception as e:
            logger.error(f"文件解析失败：{str(e)}")
            return {"error": str(e)}
    
    @staticmethod
    async def _parse_pdf(file_path: str) -> Dict[str, Any]:
        """解析 PDF 文件"""
        try:
            reader = PdfReader(file_path)
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
            
            return {
                "type": "pdf",
                "pages": len(reader.pages),
                "text": text,
                "preview": text[:2000]
            }
        except Exception as e:
            return {"error": f"PDF 解析失败：{str(e)}"}
    
    @staticmethod
    async def _parse_docx(file_path: str) -> Dict[str, Any]:
        """解析 Word 文件"""
        try:
            doc = Document(file_path)
            text = "\n".join([para.text for para in doc.paragraphs])
            
            return {
                "type": "docx",
                "paragraphs": len(doc.paragraphs),
                "text": text,
                "preview": text[:2000]
            }
        except Exception as e:
            return {"error": f"Word 解析失败：{str(e)}"}
    
    @staticmethod
    async def _parse_excel(file_path: str) -> Dict[str, Any]:
        """解析 Excel 文件"""
        try:
            df = pd.read_excel(file_path)
            
            return {
                "type": "excel",
                "rows": len(df),
                "columns": list(df.columns),
                "data": df.to_dict(orient="records"),
                "preview": df.head(10).to_dict(orient="records")
            }
        except Exception as e:
            return {"error": f"Excel 解析失败：{str(e)}"}
    
    @staticmethod
    def detect_file_type(file_path: str) -> str:
        """检测文件类型"""
        ext = file_path.split(".")[-1].lower() if "." in file_path else ""
        
        if ext in ["xlsx", "xls"]:
            return "excel"
        elif ext == "csv":
            return "csv"
        elif ext == "pdf":
            return "pdf"
        elif ext in ["docx", "doc"]:
            return "docx"
        else:
            return "unknown"


# 全局解析器实例
document_parser = DocumentParser()
